import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. Get the list of screenshots sorted chronologically by name
files = os.listdir('.')
screenshots = sorted([f for f in files if f.startswith('Screenshot') and f.endswith('.png')])

if len(screenshots) != 8:
    print(f"Expected 8 screenshots, found {len(screenshots)}.")
    # Still proceed, just use what's available

# 2. Create the Word Document
doc = Document()

# Title
title = doc.add_heading('AI-Powered College Assistant Assignment', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Define headings in order
headings = [
    "1. Source Code",
    "2. Test Case 1: Attendance Calculator",
    "3. Test Case 2: Result Calculator",
    "4. Test Case 3: Fee Balance Calculator",
    "5. Test Case 4: Library Fine Calculator",
    "6. Test Case 5: Hostel Fee Calculator",
    "7. Test Case 6: Multi-Tool Challenge",
    "8. Test Case 7: Bonus Challenge (Student Lookup)"
]

# Insert headings and corresponding images
for i in range(len(headings)):
    doc.add_heading(headings[i], level=1)
    if i < len(screenshots):
        try:
            doc.add_picture(screenshots[i], width=Inches(6.0))
        except Exception as e:
            print(f"Could not add image {screenshots[i]}: {e}")
            doc.add_paragraph(f"[Error loading {screenshots[i]}]")
    else:
        doc.add_paragraph("[Missing Screenshot]")
    
    # Add page break except for the last item
    if i < len(headings) - 1:
        doc.add_page_break()

doc.save('College_Assistant_Assignment_Final.docx')
print("Document generated successfully as College_Assistant_Assignment_Final.docx.")
